import io
import logging
from pathlib import Path

from langchain_core.documents import Document

from app.utils.errors import DocumentProcessingError, DocumentValidationError

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".epub", ".txt", ".docx"}


def parse_document(filename: str, content: bytes) -> list[Document]:
    """
    Parse uploaded file bytes into a list of LangChain Document objects.
    Each document carries page_content and metadata with the source filename.
    """
    ext = Path(filename).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise DocumentValidationError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    try:
        if ext == ".pdf":
            return _parse_pdf(filename, content)
        elif ext == ".epub":
            return _parse_epub(filename, content)
        elif ext == ".txt":
            return _parse_txt(filename, content)
        elif ext == ".docx":
            return _parse_docx(filename, content)
    except DocumentValidationError:
        raise
    except Exception:
        logger.exception("Failed to parse document '%s'", filename)
        raise DocumentProcessingError()


# ------------------------------------------------------------------ #
# Individual parsers
# ------------------------------------------------------------------ #


def _parse_pdf(filename: str, content: bytes) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    docs = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": filename, "page": page_num},
                )
            )
    logger.info("Parsed %d pages from PDF '%s'", len(docs), filename)
    return docs


def _parse_epub(filename: str, content: bytes) -> list[Document]:
    import ebooklib
    from bs4 import BeautifulSoup
    from ebooklib import epub

    book = epub.read_epub(content)
    docs = []
    chapter_num = 0

    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            chapter_num += 1
            soup = BeautifulSoup(item.get_content(), "html.parser")
            text = soup.get_text(separator="\n").strip()
            if text:
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"source": filename, "chapter": chapter_num},
                    )
                )

    logger.info("Parsed %d chapters from EPUB '%s'", len(docs), filename)
    return docs


def _parse_txt(filename: str, content: bytes) -> list[Document]:
    text = content.decode("utf-8", errors="replace").strip()
    docs = []
    if text:
        docs.append(
            Document(
                page_content=text,
                metadata={"source": filename},
            )
        )
    logger.info("Parsed TXT '%s' (%d chars)", filename, len(text))
    return docs


def _parse_docx(filename: str, content: bytes) -> list[Document]:
    from docx import Document as DocxDocument

    docx = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    docs = []
    if text:
        docs.append(
            Document(
                page_content=text,
                metadata={"source": filename},
            )
        )
    logger.info("Parsed DOCX '%s' (%d paragraphs)", filename, len(paragraphs))
    return docs
